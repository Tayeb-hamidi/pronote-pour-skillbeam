"""DOCX exporter plugin."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from shared.exporters.base import BaseExporter
from shared.exporters.branding import (
    DEFAULT_EXPORT_TITLE,
    bool_option,
    collect_choice_rows,
    label_item_type,
    resolve_skillbeam_logo,
    split_expected_answers,
)
from shared.schemas import ContentSetResponse, ExportArtifact


class DocxExporter(BaseExporter):
    format_name = "docx"

    def export(
        self, content_set: ContentSetResponse, options: dict, output_dir: Path
    ) -> ExportArtifact:
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = options.get("filename", f"content_{content_set.project_id}.docx")
        file_path = output_dir / filename

        document = Document()
        logo_path = resolve_skillbeam_logo(options)
        if logo_path is not None:
            logo_paragraph = document.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_paragraph.add_run().add_picture(str(logo_path), width=Inches(2.6))

        title = str(options.get("title", DEFAULT_EXPORT_TITLE))
        show_correct_answers = bool_option(options.get("show_correct_answers"), default=False)

        document.add_heading(title, level=1)
        intro = document.add_paragraph(
            "Support QCM lisible. Cochez les cases pour selectionner la ou les bonnes reponses."
        )
        intro.style = document.styles["Normal"]

        for index, item in enumerate(content_set.items, start=1):
            item_label = label_item_type(item.item_type.value)
            document.add_heading(f"Question {index} - {item_label}", level=2)
            document.add_paragraph(item.prompt)

            if item.item_type.value in {"mcq", "poll"}:
                helper_line = (
                    "Cochez la bonne reponse." if item.item_type.value == "mcq" else "Cochez les bonnes reponses."
                )
                document.add_paragraph(helper_line)

                rows = collect_choice_rows(
                    correct_answer=item.correct_answer,
                    distractors=item.distractors,
                    answer_options=item.answer_options,
                )
                for text, is_correct in rows:
                    checkbox = "[x]" if show_correct_answers and is_correct else "[ ]"
                    document.add_paragraph(f"{checkbox} {text}")

                if show_correct_answers:
                    answers = split_expected_answers(item.correct_answer)
                    if answers:
                        document.add_paragraph(f"Corrige: {' | '.join(answers)}")
            else:
                if show_correct_answers and item.correct_answer:
                    document.add_paragraph(f"Reponse attendue: {item.correct_answer}")
                else:
                    document.add_paragraph("Reponse: ________________________________")

        document.save(file_path)

        return ExportArtifact(
            artifact_path=str(file_path),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
